import os
import time
import json
import unicodedata

from PIL import Image
import io
import re
from html import unescape

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys

from docx import Document
from docx.shared import Inches

course_id = "ao62et"

# This is the default delay. Increase this if you are not getting images in your PDF.
delay = 0.2

# Selenium Path. Download Selenium from Chrome for your OS
selenium_path = r"C:\Programming\Python\Selenium\chromedriver.exe"

# Enable this for getting previous errors
error_log = False

# driver = webdriver.Chrome(selenium_path)
driver = None


def startup():
    global driver
    options = webdriver.ChromeOptions()
    options.add_argument("start-maximized")
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option('useAutomationExtension', False)
    settings = {"recentDestinations": [{"id": "Save as PDF", "origin": "local", "account": ""}],
                "selectedDestinationId": "Save as PDF", "version": 2}
    prefs = {'printing.print_preview_sticky_settings.appState': json.dumps(settings)}
    options.add_experimental_option('prefs', prefs)
    options.add_argument('--kiosk-printing')

    driver = webdriver.Chrome(selenium_path, options=options)

    driver.get(f"https://iam.macmillanlearning.com/login?retURL="
               f"https%3A%2F%2Fachieve.macmillanlearning.com%2Fcourses%3Fredirect%3D%252Fcourses%252F{course_id}")


def process_html(text, number, groups, doc):
    ans, s = [x.group(1) for x in re.finditer('aria-label="(.*?)"', text)], ""
    for i in range(len(ans)):
        if groups:
            s += unescape(
                f"Q{number}. {groups}\n" if i == 0 else f"  {chr(64 + i)}. {ans[i - 1]}\n"
            )
            if i == len(ans) - 1:
                s += unescape(
                    f"  {chr(65 + i)}. {ans[i]}\n"
                )
        else:
            s += unescape(
                f"Q{number}. {ans[i]}\n" if i == 0 else f"  {chr(64 + i)}. {ans[i]}\n"
            )
    doc.add_paragraph(unescape(s))


def main(doc_name, num_q):
    global delay

    document = Document()
    if "/" in doc_name:
        document.add_heading(doc_name.split("/")[-1], 0)
    else:
        document.add_heading(doc_name, 0)
    for i in range(1, int(num_q) + 1):
        try:
            time.sleep(delay)
            windows_before = driver.current_window_handle
            driver.switch_to.default_content()
            driver.switch_to.frame(driver.find_elements(By.TAG_NAME, "iframe")[0])
            closed_bar = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.ID, "side-nav"))).get_attribute('innerHTML')
            if closed_bar and "menu-selected" in closed_bar:
                WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="side-nav"]/ul/li/button'))).click()
            try:
                WebDriverWait(driver, 0.5).until(
                    EC.element_to_be_clickable(
                        (By.XPATH, f"/html/body/div/main/section/div[1]/div/button[2]"))).click()

            except:
                try:
                    WebDriverWait(driver, 0.5).until(
                        EC.element_to_be_clickable(
                            (By.XPATH, f'/html/body/div[2]/main/nav/ul/li[2]/span/ul/li[{i}]/div/button'))).click()
                except:
                    WebDriverWait(driver, 0.5).until(
                        EC.element_to_be_clickable(
                            (By.XPATH, f'/html/body/div/main/nav/ul/li[2]/span/ul/li[{i}]/div/button'))).click()

            img = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CLASS_NAME, "module-container")))
            # Wait for table or figures
            m = re.search(r'(\(Figure.*?|\(Table.*?|<img)', driver.page_source, flags=re.IGNORECASE)
            if bool(m):
                time.sleep(delay * 10)
                print("Waiting for Table / Figure")
            table_fig = m.groups()[0][1:] if m else None
            img = img.screenshot_as_png
            imageStream = io.BytesIO(img)
            im = Image.open(imageStream)
            im = im.resize([int(0.5 * s) for s in im.size])
            im.save(f"pages/page-{i}.png")
            try:
                process_html(WebDriverWait(driver, 1).until(
                    EC.element_to_be_clickable((By.CLASS_NAME, "module-graded-correct "))).get_attribute(
                    'innerHTML'), i, table_fig, document)
            except:
                process_html(WebDriverWait(driver, 1).until(
                    EC.element_to_be_clickable((By.CLASS_NAME, "module-graded-missed "))).get_attribute(
                    'innerHTML'), i, table_fig, document)
            document.add_picture(f"pages/page-{i}.png")

            if error_log:
                # Check if errors
                WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="attempts-button"]'))).click()

                # Check if 1 wrong attempt
                if "Attempt 1" in driver.page_source:
                    # Open attempt 1
                    WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="attempts"]/li[2]/button'))).click()

                    # Screenshot
                    img = WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.CLASS_NAME, "module-container"))).screenshot_as_png
                    imageStream = io.BytesIO(img)
                    im = Image.open(imageStream)
                    im = im.resize([int(0.5 * s) for s in im.size])
                    im.save(f"pages/page-{i}-wrong-{1}.png")
                    document.add_picture(f"pages/page-{i}-wrong-{1}.png")

                    # Open Menu Again
                    WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="attempts-button"]'))).click()

                if "Attempt 2" in driver.page_source:
                    # Open attempt 2
                    WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable(
                            (By.XPATH, '/html/body/div[2]/main/section/div[1]/nav/ul/li[3]/button'))).click()

                    img = WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.CLASS_NAME, "module-container"))).screenshot_as_png
                    imageStream = io.BytesIO(img)
                    im = Image.open(imageStream)
                    im = im.resize([int(0.5 * s) for s in im.size])
                    im.save(f"pages/page-{i}-wrong-{2}.png")
                    document.add_picture(f"pages/page-{i}-wrong-{2}.png")

            driver.switch_to.default_content()
        except Exception as e:
            driver.switch_to.default_content()
            print(e)
            continue
    document.save(f'{doc_name}.docx')

    # cleanup
    for item in os.listdir("pages"):
        if item.endswith(".png"):
            os.remove(os.path.join("pages", item))


if __name__ == '__main__':
    username = input("Enter Achieve Email: ")
    password = input("Enter Achieve Password: ")
    startup()
    if username is not None:
        WebDriverWait(driver, 2).until(EC.element_to_be_clickable((By.ID, "username"))).send_keys(username)
        WebDriverWait(driver, 2).until(EC.element_to_be_clickable((By.ID, "password"))).send_keys(password)
        WebDriverWait(driver, 2).until(EC.element_to_be_clickable((By.ID, "password"))).send_keys(Keys.ENTER)
    while True:
        try:
            doc_name = input("Enter Quiz Name: ")
            num_q = input("Enter number of questions: ")
            main(doc_name, num_q)
        except:
            driver.switch_to.default_content()
            os.remove(f"{doc_name}.docx")
            print("Sorry an error has occurred Please try using the other file...")
            pass
